"""
DOCX formatter for ESG analysis reports.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List
import re

class DocxFormatter:
    """Handles formatting and export of ESG analysis to DOCX format."""
    
    def __init__(self):
        self.document = Document()
        self._setup_document()
        
    def _setup_document(self):
        """Set up basic document formatting."""
        # Set default font
        style = self.document.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        # Set margins
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
    def _add_heading(self, text: str, level: int = 1):
        """Add a formatted heading."""
        heading = self.document.add_heading(text, level=level)
        heading.style.font.name = 'Arial'
        heading.style.font.size = Pt(14 if level == 1 else 12)
        heading.style.font.bold = True
        return heading
        
    def _add_table(self, rows: int, cols: int, header: List[str] = None):
        """Add a formatted table."""
        table = self.document.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        if header:
            for i, text in enumerate(header):
                cell = table.cell(0, i)
                cell.text = text
                for paragraph in cell.paragraphs:
                    paragraph.style.font.name = 'Arial'
                    paragraph.style.font.size = Pt(10)
                    paragraph.style.font.bold = True
                    
        return table
        
    def _add_risk_matrix(self, risks: List[Dict]):
        """Add a risk matrix visualization."""
        # Create a 5x5 risk matrix
        matrix = self._add_table(6, 6)
        
        # Add headers
        headers = ['', 'Very Low', 'Low', 'Medium', 'High', 'Very High']
        for i, header in enumerate(headers):
            cell = matrix.cell(0, i)
            cell.text = header
            
        # Add impact levels
        impacts = ['Very High', 'High', 'Medium', 'Low', 'Very Low']
        for i, impact in enumerate(impacts):
            cell = matrix.cell(i + 1, 0)
            cell.text = impact
            
        # Add risks to matrix
        for risk in risks:
            likelihood = risk.get('likelihood', 0)
            impact = risk.get('impact', 0)
            cell = matrix.cell(impact + 1, likelihood + 1)
            cell.text = risk.get('name', '')
            
    def format_analysis(self, analysis: Dict):
        """Format the ESG analysis into a DOCX document."""
        # Title
        title = f"{analysis['company_name']} - ESG & Impact Analysis"
        self._add_heading(title, level=1)
        
        # Executive Summary
        self._add_heading("Executive Summary", level=1)
        self.document.add_paragraph(analysis['executive_summary'])
        
        # Business Activities
        self._add_heading("Business Activities Breakdown", level=1)
        activities_table = self._add_table(
            len(analysis['business_activities']) + 1,
            4,
            ['Activity', 'Revenue Share', 'Key ESG Factors', 'Impact Alignment']
        )
        
        for i, activity in enumerate(analysis['business_activities']):
            row = activities_table.rows[i + 1]
            row.cells[0].text = activity['name']
            row.cells[1].text = f"{activity['revenue_share']}%"
            row.cells[2].text = "\n".join(activity['key_esg_factors'])
            row.cells[3].text = activity['impact_alignment']
            
        # Environmental Analysis
        self._add_heading("Environmental Analysis", level=1)
        self.document.add_paragraph(analysis['environmental_analysis'])
        
        # Climate Impact
        self._add_heading("Climate Impact Assessment", level=2)
        climate_table = self._add_table(
            5, 2,
            ['Aspect', 'Assessment']
        )
        
        climate_aspects = [
            ('Climate Solutions', 'climate_solutions'),
            ('Vulnerability', 'vulnerability'),
            ('Adaptation', 'adaptation'),
            ('Carbon Footprint', 'carbon_footprint'),
            ('Decoupling Potential', 'decoupling')
        ]
        
        for i, (label, key) in enumerate(climate_aspects):
            row = climate_table.rows[i + 1]
            row.cells[0].text = label
            row.cells[1].text = analysis['climate_impact'][key]
            
        # Social Analysis
        self._add_heading("Social Analysis", level=1)
        self.document.add_paragraph(analysis['social_analysis'])
        
        # Governance Analysis
        self._add_heading("Governance Analysis", level=1)
        self.document.add_paragraph(analysis['governance_analysis'])
        
        # Impact Thesis Alignment
        self._add_heading("Impact Thesis Alignment", level=1)
        impact_table = self._add_table(
            6, 2,
            ['Impact Area', 'Alignment Assessment']
        )
        
        impact_areas = [
            ('Local Entrepreneurship', 'local_entrepreneurship'),
            ('Decent Jobs', 'decent_jobs'),
            ('Climate Action', 'climate_action'),
            ('Gender Empowerment', 'gender_empowerment'),
            ('Resilience', 'resilience'),
            ('Overall Impact', 'overall_impact')
        ]
        
        for i, (label, key) in enumerate(impact_areas):
            row = impact_table.rows[i + 1]
            row.cells[0].text = label
            row.cells[1].text = analysis['impact_alignment'][key]
            
        # Recommendations
        self._add_heading("Recommendations", level=1)
        
        # Due Diligence Actions
        self._add_heading("Priority Due Diligence Actions", level=2)
        for action in analysis['recommendations']['due_diligence']:
            self.document.add_paragraph(action, style='List Bullet')
            
        # ESG Clauses
        self._add_heading("Suggested ESG Clauses", level=2)
        for clause in analysis['recommendations']['esg_clauses']:
            self.document.add_paragraph(clause, style='List Bullet')
            
        # KPIs
        self._add_heading("Key Performance Indicators", level=2)
        kpi_table = self._add_table(
            len(analysis['recommendations']['kpis']) + 1,
            3,
            ['KPI', 'Target', 'Frequency']
        )
        
        for i, kpi in enumerate(analysis['recommendations']['kpis']):
            row = kpi_table.rows[i + 1]
            row.cells[0].text = kpi['name']
            row.cells[1].text = kpi['target']
            row.cells[2].text = kpi['frequency']
            
    def save(self, filename: str):
        """Save the document to a file."""
        self.document.save(filename) 